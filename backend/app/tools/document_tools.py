"""Document tools — CV/cover letter parsing and generation.

These tool functions are called by the agent to parse uploaded CVs,
generate tailored CVs, and generate tailored cover letters.
"""

from __future__ import annotations

import json
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document as DocxDocument

from app.config import settings
from app.models.user_profile import UserProfile
from app.models.job import JobListing


# ---------------------------------------------------------------------------
# CV Parsing
# ---------------------------------------------------------------------------


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file."""
    text_parts: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file."""
    doc = DocxDocument(file_path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


async def tool_parse_cv(file_path: str) -> str:
    """Parse an uploaded CV (PDF or DOCX) and extract its text content.

    Args:
        file_path: Path to the uploaded CV file (.pdf or .docx).

    Returns:
        The extracted text content of the CV.
    """
    path = Path(file_path)
    if not path.exists():
        return json.dumps({"error": f"File not found: {file_path}"})

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        text = _extract_text_from_docx(file_path)
    else:
        return json.dumps(
            {"error": f"Unsupported file format: {suffix}. Use PDF or DOCX."}
        )

    if not text.strip():
        return json.dumps({"error": "Could not extract any text from the file."})

    return json.dumps({"text": text, "file_name": path.name, "format": suffix})


# ---------------------------------------------------------------------------
# CV Generation
# ---------------------------------------------------------------------------

_CV_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; margin: 40px 50px; color: #222; line-height: 1.5; font-size: 11pt; }}
  h1 {{ font-size: 22pt; margin-bottom: 2px; color: #1a1a2e; }}
  .contact {{ color: #555; font-size: 9pt; margin-bottom: 18px; }}
  h2 {{ font-size: 13pt; color: #1a1a2e; border-bottom: 1.5px solid #1a1a2e; padding-bottom: 3px; margin-top: 20px; margin-bottom: 8px; }}
  .summary {{ font-style: italic; margin-bottom: 15px; }}
  .entry {{ margin-bottom: 12px; }}
  .entry-header {{ display: flex; justify-content: space-between; font-weight: bold; }}
  .entry-sub {{ color: #555; font-size: 10pt; }}
  ul {{ margin: 4px 0 0 18px; padding: 0; }}
  li {{ margin-bottom: 2px; }}
  .skills {{ display: flex; flex-wrap: wrap; gap: 6px; }}
  .skill-tag {{ background: #e8eaf6; padding: 2px 10px; border-radius: 12px; font-size: 9pt; }}
</style>
</head>
<body>
  <h1>{name}</h1>
  <div class="contact">{contact_line}</div>

  <h2>Professional Summary</h2>
  <div class="summary">{summary}</div>

  <h2>Skills</h2>
  <div class="skills">{skills_html}</div>

  <h2>Experience</h2>
  {experience_html}

  <h2>Education</h2>
  {education_html}

  {extras_html}
</body>
</html>"""


def _render_cv_html(profile: UserProfile) -> str:
    """Render profile data into HTML for PDF conversion."""
    contact_parts = [
        p
        for p in [profile.email, profile.phone, profile.location, profile.linkedin]
        if p
    ]
    contact_line = " · ".join(contact_parts)

    all_skills = profile.technical_skills + profile.soft_skills
    skills_html = "".join(f'<span class="skill-tag">{s}</span>' for s in all_skills)

    exp_blocks: list[str] = []
    for exp in profile.experience:
        dates = ""
        if exp.start_date:
            end = exp.end_date.strftime("%b %Y") if exp.end_date else "Present"
            dates = f"{exp.start_date.strftime('%b %Y')} – {end}"
        highlights = "".join(f"<li>{h}</li>" for h in exp.highlights)
        exp_blocks.append(
            f'<div class="entry">'
            f'<div class="entry-header"><span>{exp.title}</span><span>{dates}</span></div>'
            f'<div class="entry-sub">{exp.company}{" · " + exp.location if exp.location else ""}</div>'
            f"{('<p>' + exp.description + '</p>') if exp.description else ''}"
            f"{'<ul>' + highlights + '</ul>' if highlights else ''}"
            f"</div>"
        )
    experience_html = "\n".join(exp_blocks) if exp_blocks else "<p>—</p>"

    edu_blocks: list[str] = []
    for edu in profile.education:
        dates = ""
        if edu.start_date:
            end = edu.end_date.strftime("%Y") if edu.end_date else "Present"
            dates = f"{edu.start_date.strftime('%Y')} – {end}"
        grade = f" ({edu.grade})" if edu.grade else ""
        edu_blocks.append(
            f'<div class="entry">'
            f'<div class="entry-header"><span>{edu.degree}{grade}</span><span>{dates}</span></div>'
            f'<div class="entry-sub">{edu.institution}{" · " + edu.location if edu.location else ""}</div>'
            f"</div>"
        )
    education_html = "\n".join(edu_blocks) if edu_blocks else "<p>—</p>"

    extras_parts: list[str] = []
    if profile.certifications:
        extras_parts.append(
            "<h2>Certifications</h2><ul>"
            + "".join(f"<li>{c}</li>" for c in profile.certifications)
            + "</ul>"
        )
    if profile.languages:
        extras_parts.append(
            "<h2>Languages</h2><p>" + ", ".join(profile.languages) + "</p>"
        )
    extras_html = "\n".join(extras_parts)

    return _CV_HTML_TEMPLATE.format(
        name=profile.name,
        contact_line=contact_line,
        summary=profile.summary,
        skills_html=skills_html,
        experience_html=experience_html,
        education_html=education_html,
        extras_html=extras_html,
    )


def _generate_docx(profile: UserProfile) -> str:
    """Generate a DOCX CV from the profile and return the file path."""
    doc = DocxDocument()
    doc.add_heading(profile.name, level=0)

    contact_parts = [p for p in [profile.email, profile.phone, profile.location] if p]
    if contact_parts:
        doc.add_paragraph(" · ".join(contact_parts))

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(profile.summary)

    doc.add_heading("Skills", level=1)
    all_skills = profile.technical_skills + profile.soft_skills
    doc.add_paragraph(", ".join(all_skills))

    doc.add_heading("Experience", level=1)
    for exp in profile.experience:
        dates = ""
        if exp.start_date:
            end = exp.end_date.strftime("%b %Y") if exp.end_date else "Present"
            dates = f" ({exp.start_date.strftime('%b %Y')} – {end})"
        doc.add_heading(f"{exp.title} — {exp.company}{dates}", level=2)
        if exp.description:
            doc.add_paragraph(exp.description)
        for h in exp.highlights:
            doc.add_paragraph(h, style="List Bullet")

    doc.add_heading("Education", level=1)
    for edu in profile.education:
        grade = f" ({edu.grade})" if edu.grade else ""
        doc.add_heading(f"{edu.degree}{grade} — {edu.institution}", level=2)

    out_path = str(settings.output_dir / f"cv_{uuid.uuid4().hex[:8]}.docx")
    doc.save(out_path)
    return out_path


async def tool_generate_cv(profile_json: str, job_json: str | None = None) -> str:
    """Generate a tailored CV as PDF and DOCX from the user profile.

    The CV is rendered from the profile data. If a job listing is provided,
    the LLM should have already tailored the profile data before calling this.

    Args:
        profile_json: JSON string of the UserProfile.
        job_json: Optional JSON string of the target JobListing (for reference).

    Returns:
        JSON with paths to the generated PDF and DOCX files.
    """
    try:
        profile = UserProfile.model_validate_json(profile_json)
    except Exception as e:
        return json.dumps({"error": f"Invalid profile JSON: {e}"})

    # Generate DOCX
    docx_path = _generate_docx(profile)

    # Generate PDF via WeasyPrint
    html = _render_cv_html(profile)
    pdf_path = str(settings.output_dir / f"cv_{uuid.uuid4().hex[:8]}.pdf")
    try:
        from weasyprint import HTML

        HTML(string=html).write_pdf(pdf_path)
    except Exception as e:
        return json.dumps(
            {
                "docx_path": docx_path,
                "pdf_path": None,
                "pdf_error": str(e),
            }
        )

    return json.dumps({"pdf_path": pdf_path, "docx_path": docx_path})


# ---------------------------------------------------------------------------
# LaTeX CV Generation
# ---------------------------------------------------------------------------

# Path to the LaTeX template file (project root)
_TEMPLATE_PATH = Path(__file__).resolve().parents[3] / "cv-template.tex"


def _read_template_preamble() -> str:
    """Read the template and return everything up to (but not including)
    the DOCUMENT BODY banner / \\begin{document}."""
    text = _TEMPLATE_PATH.read_text(encoding="utf-8")
    # Split just before \begin{document}
    marker = r"\begin{document}"
    idx = text.find(marker)
    if idx == -1:
        raise RuntimeError("Template is missing \\begin{document}")
    return text[:idx]


def _get_template_text() -> str:
    """Return the full template text for the LLM to reference."""
    return _TEMPLATE_PATH.read_text(encoding="utf-8")


async def tool_generate_cv_latex(latex_body: str) -> str:
    r"""Compile a LaTeX CV to PDF.

    The LLM provides the **complete** ``.tex`` file content (preamble
    configuration + \\begin{document} ... \\end{document}).  This function
    writes it to a temp directory, runs ``pdflatex`` twice (for page
    references), and copies the resulting PDF into the output directory.

    Args:
        latex_body: The full LaTeX source for the CV (complete .tex file).

    Returns:
        JSON with ``pdf_path`` on success, or ``error`` on failure.
    """
    file_id = uuid.uuid4().hex[:8]
    pdf_filename = f"cv_{file_id}.pdf"
    pdf_out = settings.output_dir / pdf_filename

    with tempfile.TemporaryDirectory() as tmpdir:
        tex_path = Path(tmpdir) / "cv.tex"
        tex_path.write_text(latex_body, encoding="utf-8")

        # Run pdflatex twice (resolves page refs / lastpage)
        for pass_num in (1, 2):
            result = subprocess.run(
                [
                    "pdflatex",
                    "-interaction=nonstopmode",
                    "-halt-on-error",
                    "-output-directory",
                    tmpdir,
                    str(tex_path),
                ],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode != 0:
                # Extract meaningful error lines from the log
                log_lines = result.stdout.splitlines()
                error_lines = [
                    line
                    for line in log_lines
                    if line.startswith("!") or "Error" in line
                ]
                error_msg = "\n".join(error_lines[:10]) or result.stdout[-2000:]
                return json.dumps(
                    {
                        "error": f"pdflatex pass {pass_num} failed",
                        "details": error_msg,
                    }
                )

        compiled_pdf = Path(tmpdir) / "cv.pdf"
        if not compiled_pdf.exists():
            return json.dumps({"error": "pdflatex produced no PDF output"})

        # Copy to output directory
        import shutil

        shutil.copy2(str(compiled_pdf), str(pdf_out))

    return json.dumps(
        {
            "pdf_path": str(pdf_out),
            "pdf_filename": pdf_filename,
            "download_url": f"/api/documents/{pdf_filename}/download",
        }
    )


# ---------------------------------------------------------------------------
# Cover Letter Generation
# ---------------------------------------------------------------------------

_COVER_LETTER_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; margin: 50px 60px; color: #222; line-height: 1.7; font-size: 11pt; }}
  .header {{ margin-bottom: 30px; }}
  .name {{ font-size: 16pt; font-weight: bold; color: #1a1a2e; }}
  .contact {{ color: #555; font-size: 9pt; }}
  .date {{ margin: 20px 0; }}
  .body-text {{ white-space: pre-wrap; }}
</style>
</head>
<body>
  <div class="header">
    <div class="name">{name}</div>
    <div class="contact">{contact_line}</div>
  </div>
  <div class="date">{date}</div>
  <div class="body-text">{body}</div>
</body>
</html>"""


async def tool_generate_cover_letter(
    profile_json: str,
    cover_letter_text: str,
    job_json: str | None = None,
) -> str:
    """Generate a cover letter as PDF and DOCX.

    The LLM should compose the cover letter text and pass it here for
    formatting and file generation.

    Args:
        profile_json: JSON string of the UserProfile.
        cover_letter_text: The full cover letter body text (written by the LLM).
        job_json: Optional JSON string of the target JobListing (for reference).

    Returns:
        JSON with paths to the generated PDF and DOCX files.
    """
    try:
        profile = UserProfile.model_validate_json(profile_json)
    except Exception as e:
        return json.dumps({"error": f"Invalid profile JSON: {e}"})

    from datetime import date as date_cls

    contact_parts = [p for p in [profile.email, profile.phone, profile.location] if p]
    contact_line = " · ".join(contact_parts)

    # DOCX
    doc = DocxDocument()
    doc.add_heading(profile.name, level=0)
    if contact_parts:
        doc.add_paragraph(" · ".join(contact_parts))
    doc.add_paragraph(date_cls.today().strftime("%d %B %Y"))
    doc.add_paragraph("")
    doc.add_paragraph(cover_letter_text)

    docx_path = str(settings.output_dir / f"cover_letter_{uuid.uuid4().hex[:8]}.docx")
    doc.save(docx_path)

    # PDF
    html = _COVER_LETTER_HTML.format(
        name=profile.name,
        contact_line=contact_line,
        date=date_cls.today().strftime("%d %B %Y"),
        body=cover_letter_text,
    )
    pdf_path = str(settings.output_dir / f"cover_letter_{uuid.uuid4().hex[:8]}.pdf")
    try:
        from weasyprint import HTML

        HTML(string=html).write_pdf(pdf_path)
    except Exception as e:
        return json.dumps(
            {
                "docx_path": docx_path,
                "pdf_path": None,
                "pdf_error": str(e),
            }
        )

    return json.dumps({"pdf_path": pdf_path, "docx_path": docx_path})
